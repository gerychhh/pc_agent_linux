from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any


def _result(ok: bool, **kwargs: Any) -> str:
    payload = {"ok": ok, **kwargs}
    if not ok and "error" not in payload:
        payload["error"] = "unknown_error"
    return json.dumps(payload, ensure_ascii=False)


def read_file(path: str, max_chars: int = 20000) -> str:
    try:
        target = Path(path)
        content = target.read_text(encoding="utf-8")
        if len(content) > max_chars:
            content = content[:max_chars]
        return _result(True, content=content)
    except Exception as exc:
        return _result(False, error=str(exc))


def write_file(path: str, content: str) -> str:
    try:
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content, encoding="utf-8")
        return _result(True, path=str(target))
    except Exception as exc:
        return _result(False, error=str(exc))


def get_known_paths() -> str:
    home = os.path.expanduser("~")
    desktop = os.path.join(home, "Desktop")
    if not os.path.isdir(desktop):
        fallback_home = os.environ.get("USERPROFILE") or home
        desktop = os.path.join(fallback_home, "Desktop")
    documents = os.path.join(home, "Documents")
    downloads = os.path.join(home, "Downloads")
    return _result(
        True,
        desktop=desktop,
        documents=documents,
        downloads=downloads,
        home=home,
    )


def write_text_file_lines(
    path: str,
    line_template: str,
    count: int,
    add_newline: bool = True,
) -> str:
    try:
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        with target.open("w", encoding="utf-8") as handle:
            for i in range(int(count)):
                line = line_template.format(i=i, n=i + 1)
                if add_newline:
                    line += "\n"
                handle.write(line)
        return _result(True, path=str(target), written_lines=int(count), error=None)
    except Exception as exc:
        return _result(False, path=path, written_lines=0, error=str(exc))


def create_docx(path: str, title: str | None, paragraphs: list[str]) -> str:
    try:
        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            return _result(False, path=path, paragraphs=0, error=f"missing_python_docx: {exc}")
        target = Path(path)
        target.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        if title:
            document.add_heading(title, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        document.save(str(target))
        return _result(True, path=str(target), paragraphs=len(paragraphs), error=None)
    except Exception as exc:
        return _result(False, path=path, paragraphs=0, error=str(exc))
